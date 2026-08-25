#!/usr/bin/env python3
"""Gera o Relatório Final do Lab01 a partir do template da disciplina (#11)."""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "enunciados" / "Template_Relatorio_Laboratorio (1).docx"
OUT = ROOT / "relatorio" / "Lab01_Relatorio_Final.docx"


def set_run_normal(run) -> None:
    run.italic = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    try:
        run.font.highlight_color = None
    except Exception:
        pass


def replace_paragraph_text(paragraph, text: str, *, bold_first_sentence: bool = False) -> None:
    """Substitui todo o texto do parágrafo mantendo o estilo do parágrafo."""
    # limpa runs
    for _ in range(len(paragraph.runs)):
        paragraph.runs[0].clear()
        r = paragraph.runs[0]._element
        r.getparent().remove(r)
    # se ainda sobrou texto no elemento
    if paragraph.text:
        paragraph.clear()

    if not text:
        return

    run = paragraph.add_run(text)
    set_run_normal(run)
    run.font.size = Pt(11)


def fill_cell(cell, text: str) -> None:
    # limpa e escreve na primeira paragraph
    for p in cell.paragraphs:
        for _ in range(len(p.runs)):
            p.runs[0].clear()
            el = p.runs[0]._element
            el.getparent().remove(el)
    if not cell.paragraphs:
        cell.add_paragraph(text)
    else:
        cell.paragraphs[0].add_run(text)


def main() -> None:
    if not TEMPLATE.exists():
        raise SystemExit(f"Template não encontrado: {TEMPLATE}")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(TEMPLATE, OUT)
    doc = Document(str(OUT))

    # --- Capa (tabela 0) ---
    capa = doc.tables[0]
    # rows: Curso, Disciplina, Turno, Professor, Laboratorio, Grupo, Link, Data
    fill_cell(capa.rows[4].cells[1], "Lab01 — Características de repositórios populares + Setup do Kanban")
    fill_cell(
        capa.rows[5].cells[1],
        "Gabriel Ferreira Amaral (Druitti) · Gabriel Santiago (gabsant07) · Gabriel Tavares (Tavaresds1)",
    )
    fill_cell(
        capa.rows[6].cells[1],
        "Repo: https://github.com/Druitti/Laboratorio-de-Desenvolvimento-de-Software | "
        "Project: https://github.com/users/Druitti/projects/5",
    )
    fill_cell(capa.rows[7].cells[1], "25/08/2026")

    # Mapa de parágrafos a substituir (índices do template)
    # Remove ORIENTAÇÃO e placeholders
    content = {
        2: (
            "Relatório Final do Laboratório 01 da disciplina Laboratório de Experimentação de Software "
            "(PUC Minas). Este documento descreve a mineração dos 1.000 repositórios mais estrelados do "
            "GitHub, as hipóteses informais do grupo, a metodologia de coleta via GraphQL e a configuração "
            "do processo Kanban no GitHub Projects. A seção de visualização gráfica detalhada (Issue #10) "
            "é de responsabilidade dos demais integrantes e deve ser anexada/integrada quando concluída."
        ),
        6: "",  # remove orientação intro
        12: (
            "Repositórios open-source populares concentram boa parte da atenção da indústria e da academia, "
            "mas nem sempre há evidência sistemática sobre o perfil desses sistemas (idade, contribuição "
            "externa, releases, atualização, linguagem e gestão de issues). Este laboratório investiga essas "
            "características nos 1.000 repositórios com maior número de estrelas no GitHub e, em paralelo, "
            "institui o quadro Kanban (GitHub Projects v2) que acompanhará o grupo no semestre.\n\n"
            "Questões de Pesquisa do enunciado:\n"
            "RQ01. Sistemas populares são maduros/antigos? (idade a partir de createdAt)\n"
            "RQ02. Sistemas populares recebem muita contribuição externa? (PRs merged)\n"
            "RQ03. Sistemas populares lançam releases com frequência? (total de releases)\n"
            "RQ04. Sistemas populares são atualizados com frequência? (tempo desde pushedAt)\n"
            "RQ05. Sistemas populares são escritos nas linguagens mais populares? (primaryLanguage vs TIOBE Index)\n"
            "RQ06. Sistemas populares possuem alto percentual de issues fechadas? (closed/total)\n"
            "RQ07 (bônus). Linguagens TIOBE associam-se a mais PRs, releases e atualizações mais recentes?\n\n"
            "Hipóteses informais (antes da análise completa — Issue #9):\n"
            "RQ01 — Sim: mediana de idade elevada (vários anos), pois popularidade acumula-se no tempo.\n"
            "RQ02 — Em parte: mediana de PRs merged relevante, com forte assimetria (poucos repos concentram milhares de PRs).\n"
            "RQ03 — Não necessariamente: mediana de releases baixa/próxima de zero por causa de repos de curadoria/docs.\n"
            "RQ04 — Sim: tempo desde o último push tipicamente curto (dias/poucas semanas).\n"
            "RQ05 — Parcialmente: parcela relevante em linguagens TIOBE, mas também Markdown/null e linguagens fora do índice (ex.: TypeScript).\n"
            "RQ06 — Sim: razão mediana closed/total alta (~70–80% ou mais).\n"
            "RQ07 — Mista: TIOBE com mais PRs e updates mais recentes; releases com diferença menor.\n\n"
            "Inovações (30%): (i) validação rápida de amostra (5–10 repos) por bloco de métricas antes da coleta completa; "
            "(ii) paginação em lotes menores com retry/backoff para mitigar HTTP 502 na API GraphQL; "
            "(iii) snapshots CSV do GitHub Projects para reconstituição histórica do board (base dos Labs 04/05)."
        ),
        14: "",
        15: (
            "Este é o Lab01, ponto de partida do semestre: combina mineração de software (1.000 repositórios "
            "mais populares do GitHub) com a implantação do processo Kanban do próprio grupo. Os dados e o "
            "board gerados aqui alimentam laboratórios posteriores (em especial snapshots para Labs 04 e 05).\n\n"
            "Objeto de estudo: repositórios retornados pela busca GraphQL "
            "`stars:>1 sort:stars-desc`, limitados aos 1.000 primeiros, com métricas alinhadas às RQs 01–06 "
            "(e campos necessários ao bônus RQ07).\n\n"
            "Referência conceitual para “linguagens mais populares” (RQ05/RQ07): TIOBE Index "
            "(https://www.tiobe.com/tiobe-index/), mantida como fonte única em todo o Lab01, conforme "
            "documentado no README e em docs/hipoteses-informais.md."
        ),
        17: "",
        19: "",
        20: (
            "Principais desafios enfrentados:\n"
            "• Complexidade/instabilidade da API GraphQL do GitHub ao solicitar connections (PRs, issues, releases) "
            "para dezenas de repositórios na mesma página — ocorrência recorrente de HTTP 502, mitigada com "
            "páginas menores e retries com backoff.\n"
            "• Paginação até 1.000 itens respeitando o limite de 100 nós por página da API Search.\n"
            "• Ausência de histórico consultável de mudança de Status no GitHub Projects via API, exigindo "
            "exportações periódicas (snapshots CSV) do board.\n"
            "• Valores ausentes: primaryLanguage nulo (listas/docs) e repositórios com zero issues "
            "(razão RQ06 indefinida).\n"
            "• Restrição do enunciado: proibição de bibliotecas de terceiros que encapsulem a API do GitHub — "
            "a query GraphQL e o cliente HTTP são próprios do grupo (Python + requests)."
        ),
        22: "",
        23: (
            "Tomadas de decisão:\n"
            "• Linguagem/script: Python 3 com requests e query GraphQL própria (sem SDK GitHub).\n"
            "• Amostra: top 1.000 por estrelas via search GraphQL; sem exclusão a priori por linguagem.\n"
            "• RQ04 operacionalizada como dias desde pushedAt (atividade de código); updatedAt fica como apoio.\n"
            "• RQ05: classificação TIOBE vs. não-TIOBE a partir de primaryLanguage.name (TypeScript, Markdown etc. "
            "contam como fora do conjunto TIOBE adotado).\n"
            "• WIP da coluna Doing = 2: trio; evita excesso de paralelismo e força Review/Done frequentes "
            "(docs/processo-kanban.md).\n"
            "• Commits devem referenciar #Issue para vinculação no board (correção centrada no Projects)."
        ),
        25: "",
        26: (
            "Etapas (sprints) e responsáveis (refletindo Assignees das Issues):\n\n"
            "Lab01S01 — Coleta GraphQL (100 repos) + setup Kanban\n"
            "• #2 Setup Projects/WIP/labels — Gabriel Ferreira (Druitti) — Done\n"
            "• #1 Query base + auth + fetch — Druitti — Done\n"
            "• #3 RQ01+RQ02 + validação amostra — Druitti — Done\n"
            "• #4 RQ03+RQ04 + validação — Druitti — Done\n"
            "• #5 RQ05+RQ06 + TIOBE — Druitti — Done\n"
            "• #6 Integração/smoke 100 — Druitti — Done\n\n"
            "Lab01S02 — Paginação 1000 + CSV + hipóteses + snapshot\n"
            "• #7 Paginação 1000 + export CSV — Gabriel Santiago (gabsant07) — Done (código + dados gerados)\n"
            "• #8 Snapshot GraphQL do Project → CSV — gabsant07 — em aberto no board (código presente: "
            "src/collect_repos_snapshot.py)\n"
            "• #9 Hipóteses informais — Druitti — Done\n\n"
            "Lab01S03 — Análise/visualização\n"
            "• #10 Análise/visualização RQs + bônus — Gabriel Tavares (Tavaresds1) e gabsant07 — em andamento "
            "(fora do escopo deste responsável pelo relatório textual)\n\n"
            "Relatório Final\n"
            "• #11 Relatório final + print do board — Druitti — este documento\n\n"
            "Configuração do processo (Kanban):\n"
            "• Colunas Status: Backlog → To Do → Doing → Review → Done\n"
            "• WIP Doing = 2 (justificativa na seção 3.2 e em docs/processo-kanban.md)\n"
            "• Cartões = Issues reais com Assignee; labels por integrante\n\n"
            "[INSERIR AQUI O PRINT DO BOARD] — Anexo obrigatório: captura de tela do GitHub Projects "
            "(https://github.com/users/Druitti/projects/5) ao final do Lab01, mostrando o fluxo completo "
            "e a política de WIP. Salvar também em relatorio/anexos/board-lab01.png se desejar versionar a imagem."
        ),
        27: (
            "Sugestão de anexo: print do quadro Kanban (GitHub Projects) mencionado acima. "
            "Enquanto a imagem não for colada neste .docx, use o link do Project e/ou o arquivo em relatorio/anexos/."
        ),
        29: "",
        30: (
            "Ferramentas utilizadas:\n"
            "• GitHub GraphQL API (https://api.github.com/graphql) — mineração e snapshot do Project\n"
            "• Python 3 + requests (+ python-dotenv opcional) — scripts próprios do grupo "
            "(src/collect_repos.py, src/validate_sample.py, src/collect_repos_snapshot.py)\n"
            "• Git / GitHub Issues + GitHub Projects (v2) — processo\n"
            "• (Pendente #10) biblioteca/ferramenta de visualização escolhida pelos responsáveis pela análise\n\n"
            "Links:\n"
            "• Repositório: https://github.com/Druitti/Laboratorio-de-Desenvolvimento-de-Software\n"
            "• GitHub Projects: https://github.com/users/Druitti/projects/5"
        ),
        35: "",
        36: (
            "Inovações propostas pelo grupo (30%):\n"
            "1) Pipeline de validação de amostra (validate_sample.py) por bloco de RQs antes da coleta em escala — "
            "reduz risco de métrica ausente/errada.\n"
            "2) Estratégia anti-502 (page size reduzido + retries) na coleta GraphQL dos top repositórios.\n"
            "3) Exportação de snapshots do GitHub Projects em CSV (collect_repos_snapshot.py), criando série "
            "temporal do Status das Issues para Labs 04/05.\n\n"
            "Os resultados numéricos dessas inovações aparecem na robustez da coleta (1.000 repos obtidos) e "
            "na rastreabilidade do processo; gráficos comparativos detalhados ficam a cargo da Issue #10."
        ),
        40: (
            "Coleta de dados (estado atual):\n"
            "• Alvo: 1.000 repositórios; obtido: 1.000 linhas em data/top_1000_repos.csv "
            "(e JSON correspondente).\n"
            "• Período de referência da coleta local disponível: agosto/2026.\n"
            "• Ausências observadas na amostra: primaryLanguage nulo em 87 repos; 43 repos com zero issues "
            "(razão RQ06 não definida).\n"
            "• Tratamento: ausências reportadas; razões RQ06 calculadas apenas quando issues > 0; "
            "linguagem nula tratada como categoria própria / fora do TIOBE.\n\n"
            "Resumo descritivo (medianas na amostra de 1.000 — apoio ao relatório; gráficos na #10):\n"
            "• RQ01 — mediana da idade ≈ 2.831 dias (~7,8 anos)\n"
            "• RQ02 — mediana de PRs merged ≈ 768 (média ≈ 4.243 — forte assimetria)\n"
            "• RQ03 — mediana de releases ≈ 40; cerca de 28% com zero releases\n"
            "• RQ04 — mediana de dias desde pushedAt ≈ 6\n"
            "• RQ05 — entre repos com linguagem: ~67,9% em linguagens do conjunto TIOBE adotado; "
            "tops: Python, TypeScript, JavaScript, Go, Rust\n"
            "• RQ06 — mediana closed/total ≈ 0,875 (87,5%)\n\n"
            "Snapshots do Project: script disponível (src/collect_repos_snapshot.py); a Issue #8 permanece "
            "aberta no board até a exportação CSV de fechamento ser confirmada pelos responsáveis."
        ),
        42: "",
        44: (
            "[PENDENTE — Issue #10] Inserir aqui os gráficos (um por RQ + bônus), cada um precedido da pergunta "
            "correspondente, com eixos nomeados e medianas/valores-chave explicitados no texto. "
            "Responsáveis: Gabriel Tavares (Tavaresds1) e Gabriel Santiago (gabsant07)."
        ),
        47: "",
        48: (
            "Discussão preliminar hipótese × evidência descritiva (medianas). A discussão completa com "
            "visualizações e eventuais testes deve ser consolidada após a Issue #10.\n\n"
            "RQ01 — Hipótese confirmada preliminarmente: mediana ~7,8 anos indica maturidade elevada.\n"
            "RQ02 — Hipótese parcialmente confirmada: mediana 768 PRs merged é alta, mas a média muito superior "
            "evidencia assimetria (cauda longa).\n"
            "RQ03 — Hipótese parcialmente confirmada: há mediana > 0 (40), porém 28% sem releases — "
            "compatível com mistura de software versionado e curadoria.\n"
            "RQ04 — Hipótese confirmada preliminarmente: mediana de 6 dias desde o último push indica "
            "atualização frequente no topo de popularidade.\n"
            "RQ05 — Hipótese parcialmente confirmada: maioria (entre os com linguagem) no TIOBE, mas "
            "TypeScript/Markdown/null mostram que popularidade ≠ apenas top TIOBE.\n"
            "RQ06 — Hipótese confirmada preliminarmente: mediana 87,5% de issues fechadas.\n"
            "RQ07 — Aguardando cruzamentos por linguagem da Issue #10 para confirmar/refutar a hipótese mista.\n\n"
            "Ameaça à validade: Search API ordena por estrelas e pode omitir/alterar ranking ao longo do tempo; "
            "primaryLanguage nulo e repos sem issues afetam RQ05/RQ06; PRs merged medem contribuição aceita, "
            "não necessariamente “externa” no sentido estrito (contribuições de maintainers também entram)."
        ),
        51: (
            "O Lab01 permitiu caracterizar, de forma exploratória, o perfil dos 1.000 repositórios mais "
            "estrelados do GitHub e institucionalizar o Kanban do grupo (colunas Backlog→Done, WIP=2, Issues "
            "com Assignee e commits referenciando #N).\n\n"
            "Em síntese descritiva: os sistemas populares tendem a ser maduros e atualizados com frequência, "
            "com razão alta de issues fechadas e contribuição via PRs assimétrica; a relação com linguagens "
            "TIOBE é parcial; releases variam bastante.\n\n"
            "Limitações: dependência da Search API; métricas proxy (ex.: PRs merged ≠ contribuição externa pura); "
            "visualizações finais e aprofundamento estatístico pendentes da Issue #10; print do board deve ser "
            "anexado neste documento.\n\n"
            "Com mais tempo, o grupo expandiria a RQ07 com testes por linguagem e automatizaria snapshots a "
            "cada fechamento de sprint no CI."
        ),
    }

    for idx, text in content.items():
        if idx < len(doc.paragraphs):
            # Remove orientação: se text=="" e começa com ORIENTAÇÃO, limpa
            p = doc.paragraphs[idx]
            if text == "" and p.text.strip().startswith("ORIENTAÇÃO"):
                replace_paragraph_text(p, "")
            elif text:
                replace_paragraph_text(p, text)

    # Remover listas "Perguntas que esta seção..." da intro (deixar mais limpo) — opcional
    # Preencher tabela de métricas (table 1)
    metrics = doc.tables[1]
    rows_data = [
        ("RQ01", "Idade do repositório", "Data da coleta − createdAt", "Dias", "Script GraphQL próprio (API GitHub)"),
        ("RQ02", "PRs aceitas", "pullRequests(states:MERGED).totalCount", "Contagem", "Script GraphQL próprio"),
        ("RQ03", "Releases", "releases.totalCount", "Contagem", "Script GraphQL próprio"),
        ("RQ04", "Tempo até última atualização", "Data da coleta − pushedAt", "Dias", "Script GraphQL próprio"),
        ("RQ05", "Linguagem primária", "primaryLanguage.name classificada vs TIOBE Index", "Categoria", "GraphQL + TIOBE Index"),
        ("RQ06", "% issues fechadas", "closedIssues.totalCount / issues.totalCount (se total>0)", "Proporção", "Script GraphQL próprio"),
        ("RQ07 (bônus)", "PRs/releases/atualização por linguagem", "Segmentação das métricas RQ02–04 por linguagem / TIOBE", "Múltipla", "Pós-processamento (#10)"),
    ]
    # table has header + example + empty rows — rewrite from row 1
    while len(metrics.rows) < len(rows_data) + 1:
        metrics.add_row()
    for i, row_vals in enumerate(rows_data, start=1):
        for j, val in enumerate(row_vals):
            fill_cell(metrics.rows[i].cells[j], val)

    # Referências — acrescentar após a existente
    # Encontrar parágrafo de referências list
    refs_extra = [
        "TIOBE Software BV. TIOBE Index. Disponível em: https://www.tiobe.com/tiobe-index/. Acesso em: 25 ago. 2026.",
        "GitHub. GraphQL API documentation. Disponível em: https://docs.github.com/en/graphql. Acesso em: 25 ago. 2026.",
        "BASILI, Victor R.; CALDIERA, Gianluigi; ROMBACH, H. Dieter. The Goal Question Metric Approach. 1994.",
    ]
    # append at end
    doc.add_paragraph(
        "Grupo Lab01 — repositório e board: "
        "https://github.com/Druitti/Laboratorio-de-Desenvolvimento-de-Software ; "
        "https://github.com/users/Druitti/projects/5"
    )
    for ref in refs_extra:
        doc.add_paragraph(ref, style="List Paragraph")

    doc.save(str(OUT))
    print(f"Relatório gerado: {OUT}")


if __name__ == "__main__":
    main()
